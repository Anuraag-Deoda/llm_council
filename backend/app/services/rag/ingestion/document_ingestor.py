"""
Document ingestor for uploaded files (PDF, DOCX, TXT, MD).
"""
import os
import logging
from datetime import datetime
from pathlib import Path
from typing import List, Optional
import chardet

from .base_ingestor import BaseIngestor, IngestedDocument

logger = logging.getLogger(__name__)


class DocumentIngestor(BaseIngestor):
    """
    Ingestor for uploaded document files.

    Supports:
    - PDF files (.pdf)
    - Word documents (.docx, .doc)
    - Text files (.txt)
    - Markdown files (.md, .markdown)
    """

    SUPPORTED_TYPES = {
        'pdf': ['.pdf'],
        'docx': ['.docx', '.doc'],
        'txt': ['.txt'],
        'md': ['.md', '.markdown'],
    }

    def __init__(self, source_name: str = "uploads", base_trust_score: float = 0.8):
        super().__init__(source_name, base_trust_score)

    def get_supported_types(self) -> List[str]:
        """Get list of supported file extensions."""
        extensions = []
        for ext_list in self.SUPPORTED_TYPES.values():
            extensions.extend(ext_list)
        return extensions

    async def extract_content(
        self,
        source_path: str,
        title: Optional[str] = None,
        author: Optional[str] = None,
        **kwargs
    ) -> IngestedDocument:
        """
        Extract content from a document file.

        Args:
            source_path: Path to the document file
            title: Optional title (defaults to filename)
            author: Optional author name
            **kwargs: Additional parameters

        Returns:
            IngestedDocument with extracted content
        """
        path = Path(source_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {source_path}")

        extension = path.suffix.lower()
        file_type = self._get_file_type(extension)

        if file_type is None:
            raise ValueError(f"Unsupported file type: {extension}")

        # Extract content based on file type
        if file_type == 'pdf':
            content = await self._extract_pdf(path)
        elif file_type == 'docx':
            content = await self._extract_docx(path)
        elif file_type in ('txt', 'md'):
            content = await self._extract_text(path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        # Get file metadata
        stat = path.stat()
        created_at = datetime.fromtimestamp(stat.st_ctime)
        updated_at = datetime.fromtimestamp(stat.st_mtime)

        # Compute content hash
        content_hash = self.compute_content_hash(content)

        # Determine title
        if not title:
            title = path.stem  # Filename without extension

        return IngestedDocument(
            title=title,
            content=content,
            content_hash=content_hash,
            file_type=file_type,
            file_path=str(path.absolute()),
            author=author,
            author_trust_score=kwargs.get('author_trust_score', 0.5),
            source_created_at=created_at,
            source_updated_at=updated_at,
            extra_data={
                'file_size': stat.st_size,
                'extension': extension,
            }
        )

    async def extract_batch(
        self, source_paths: List[str], **kwargs
    ) -> List[IngestedDocument]:
        """
        Extract content from multiple document files.

        Args:
            source_paths: List of file paths
            **kwargs: Additional parameters passed to extract_content

        Returns:
            List of IngestedDocuments
        """
        documents = []
        for path in source_paths:
            try:
                doc = await self.extract_content(path, **kwargs)
                documents.append(doc)
            except Exception as e:
                logger.error(f"Failed to extract content from {path}: {e}")
                # Continue with other files
        return documents

    def _get_file_type(self, extension: str) -> Optional[str]:
        """Map file extension to file type."""
        for file_type, extensions in self.SUPPORTED_TYPES.items():
            if extension in extensions:
                return file_type
        return None

    async def _extract_pdf(self, path: Path) -> str:
        """Extract text from PDF file."""
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(str(path))
            text_parts = []

            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                except Exception as e:
                    logger.warning(f"Failed to extract page {page_num} from {path}: {e}")

            content = '\n\n'.join(text_parts)

            if not content.strip():
                logger.warning(f"No text content extracted from PDF: {path}")
                return ""

            return content

        except ImportError:
            raise ImportError("PyPDF2 is required for PDF processing. Install with: pip install PyPDF2")

    async def _extract_docx(self, path: Path) -> str:
        """Extract text from Word document."""
        try:
            from docx import Document

            doc = Document(str(path))
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))

            return '\n\n'.join(text_parts)

        except ImportError:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")

    async def _extract_text(self, path: Path) -> str:
        """Extract text from plain text or markdown file."""
        # Detect encoding
        with open(path, 'rb') as f:
            raw_data = f.read()
            detected = chardet.detect(raw_data)
            encoding = detected.get('encoding', 'utf-8') or 'utf-8'

        # Read with detected encoding
        try:
            with open(path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            # Fallback to utf-8 with error handling
            with open(path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()

    async def extract_from_bytes(
        self,
        content: bytes,
        filename: str,
        title: Optional[str] = None,
        author: Optional[str] = None,
        **kwargs
    ) -> IngestedDocument:
        """
        Extract content from file bytes (for uploaded files).

        Args:
            content: File content as bytes
            filename: Original filename
            title: Optional title
            author: Optional author
            **kwargs: Additional parameters

        Returns:
            IngestedDocument with extracted content
        """
        import tempfile

        extension = Path(filename).suffix.lower()
        file_type = self._get_file_type(extension)

        if file_type is None:
            raise ValueError(f"Unsupported file type: {extension}")

        # Write to temporary file for processing
        with tempfile.NamedTemporaryFile(suffix=extension, delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            doc = await self.extract_content(
                tmp_path,
                title=title or Path(filename).stem,
                author=author,
                **kwargs
            )
            # Update file path to indicate it was from upload
            doc.file_path = filename
            return doc
        finally:
            # Clean up temp file
            try:
                os.unlink(tmp_path)
            except Exception:
                pass
